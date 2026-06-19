#!/usr/bin/env python3
"""
verify.py — Check if references are real, find duplicates, suggest replacements.

Usage:
  python verify.py list < "refs.txt"
  python verify.py docx <path/to/thesis.docx>
  python verify.py doi 10.1000/example
  python verify.py pmid 12345678
"""
import argparse, json, os, re, sys, time
from collections import defaultdict

INSTALLED = {"requests": False, "biopython": False, "semanticscholar": False, "docx": False}

try:
    import requests
    INSTALLED["requests"] = True
except ImportError:
    pass

try:
    from Bio import Entrez
    INSTALLED["biopython"] = True
    Entrez.email = "refcheck@local"
except ImportError:
    pass

try:
    from docx import Document
    INSTALLED["docx"] = True
except ImportError:
    pass

# ── helpers ──────────────────────────────────────────────

def extract_doi(text):
    m = re.search(r'(10\.\d{4,}/[^\s,;)]+)', text)
    return m.group(1) if m else None

def extract_pmid(text):
    m = re.search(r'\b(\d{8})\b', text)
    return m.group(1) if m else None

def clean(t):
    return re.sub(r'\s+', ' ', t).strip()

# ── verification sources ─────────────────────────────────

CACHE = {}

def check_crossref(doi):
    if doi in CACHE: return CACHE[doi]
    if not INSTALLED["requests"]: return None
    try:
        r = requests.get(f"https://api.crossref.org/works/{doi}",
                         headers={"User-Agent": "RefCheck/1.0"}, timeout=8)
        if r.status_code == 200:
            d = r.json()["message"]
            authors = [a.get("family","") for a in d.get("author",[])][:5]
            result = {"found": True, "source": "CrossRef",
                "title": (d.get("title") or [""])[0],
                "authors": ", ".join(authors),
                "year": (d.get("published-print") or d.get("issued") or {}).get("date-parts",[[None]])[0][0],
                "journal": (d.get("container-title") or [""])[0], "doi": doi}
            CACHE[doi] = result
            return result
    except: pass
    CACHE[doi] = None
    return None

def check_pubmed(pmid=None, doi=None, title=None):
    if pmid and pmid in CACHE: return CACHE[pmid]
    if not INSTALLED["biopython"]: return None
    try:
        ids = []
        if pmid: ids = [pmid]
        elif doi:
            h = Entrez.esearch(db="pubmed", term=f"({doi}[DOI])", retmax=1, retmode="json")
            ids = json.load(h)["esearchresult"]["idlist"]; h.close()
        elif title:
            h = Entrez.esearch(db="pubmed", term=title, retmax=3, retmode="json")
            ids = json.load(h)["esearchresult"]["idlist"]; h.close()
        if not ids: return None
        h = Entrez.esummary(db="pubmed", id=",".join(ids[:3]), retmode="json")
        data = json.load(h); h.close()
        for uid in ids[:3]:
            rec = data["result"].get(uid)
            if rec:
                doi_val = None
                eloc = rec.get("elocationid","")
                if "doi:" in eloc: doi_val = eloc.replace("doi: ","")
                result = {"found": True, "source": "PubMed", "pmid": uid,
                    "title": rec.get("title",""),
                    "authors": rec.get("authors","").split(", ")[:3],
                    "year": rec.get("pubdate","")[:4],
                    "journal": rec.get("source",""), "doi": doi_val}
                if pmid: CACHE[pmid] = result
                return result
    except: pass
    return None

def check_semantic(title=None, doi=None):
    if not INSTALLED["requests"]: return None
    try:
        query = doi or title
        if not query: return None
        r = requests.get("https://api.semanticscholar.org/graph/v1/paper/search",
            params={"query": query, "limit": 3, "fields": "title,year,externalIds,authors,url"},
            timeout=8)
        if r.status_code == 429:
            time.sleep(1)
            r = requests.get("https://api.semanticscholar.org/graph/v1/paper/search",
                params={"query": query, "limit": 3, "fields": "title,year,externalIds,authors,url"},
                timeout=8)
        if r.status_code == 200:
            for paper in r.json().get("data", []):
                ext_ids = paper.get("externalIds", {}) or {}
                if doi and ext_ids.get("DOI","") != doi: continue
                return {"found": True, "source": "Semantic Scholar",
                    "title": paper.get("title",""),
                    "authors": [a.get("name","") for a in (paper.get("authors") or [])[:5]],
                    "year": paper.get("year"),
                    "doi": ext_ids.get("DOI",""),
                    "url": f"https://www.semanticscholar.org/paper/{paper.get('paperId','')}"}
    except: pass
    return None

def verify_ref(ref):
    doi = ref.get("doi") or extract_doi(ref.get("raw",""))
    pmid = ref.get("pmid")
    title = ref.get("title","")
    result = ref.copy()
    if doi:
        cr = check_crossref(doi)
        if cr: result.update(cr); return result
        pm = check_pubmed(doi=doi)
        if pm: result.update(pm); return result
    if pmid:
        pm = check_pubmed(pmid=pmid)
        if pm: result.update(pm); return result
    if title:
        pm = check_pubmed(title=title)
        if pm: result.update(pm); return result
        ss = check_semantic(title=title)
        if ss and ss.get("found"): result.update(ss); return result
    result.update({"found": False, "source": "none"})
    return result

# ── suggest replacement ──────────────────────────────────

STOPWORDS = {'in','of','the','and','for','with','on','from','by','to','a','an','is','was','are','were','has','have','been','that','this','these','those','their','its','our','your','his','her','not','but','or','as','at','be','we','it','et','al','vol','no','pp','doi','pmid','pmc','via'}

def extract_keywords(text):
    if not text: return ""
    t = text
    # Remove DOIs, URLs, IDs
    t = re.sub(r'10\.\S+|https?://\S+|PMID:\s*\d+|PMCID:\s*\S+', '', t, flags=re.I)
    # Years
    t = re.sub(r'\(?(?:19|20)\d{2}[a-z]?\)?', ' ', t)
    # Author names (Surname, I. or Surname I. or Surname et al.)
    t = re.sub(r'[A-Z][a-z]+,\s*(?:[A-Z]\.?\s*)+', '', t)
    t = re.sub(r'[A-Z][a-z]+\s+(?:[A-Z]\.\s*)+', '', t)
    t = re.sub(r'\bet\s+al\.?\b', '', t, flags=re.I)
    # Single letters, numbers, punctuation
    t = re.sub(r'\b[a-zA-Z]\b', ' ', t)
    t = re.sub(r'[^a-zA-Z\s]', ' ', t)
    t = re.sub(r'\s+', ' ', t).strip()
    # Filter
    words = [w for w in t.split() if w.lower() not in STOPWORDS and len(w) > 3]
    return " ".join(words[:10])

def suggest_replacement(raw_text):
    keywords = extract_keywords(raw_text)
    if not keywords or len(keywords) < 5:
        return None
    try:
        if INSTALLED["biopython"]:
            h = Entrez.esearch(db="pubmed", term=keywords, retmax=3, retmode="json", sort="relevance")
            ids = json.load(h)["esearchresult"]["idlist"]; h.close()
            if ids:
                h = Entrez.esummary(db="pubmed", id=",".join(ids[:3]), retmode="json")
                data = json.load(h); h.close()
                for uid in ids[:3]:
                    rec = data["result"].get(uid)
                    if rec and rec.get("title"):
                        authors_raw = rec.get("authors",[])
                        if isinstance(authors_raw, list):
                            authors = [a["name"] if isinstance(a, dict) else str(a) for a in authors_raw]
                        else:
                            authors = str(authors_raw).split(", ")
                        doi = None
                        eloc = rec.get("elocationid","")
                        if isinstance(eloc, str) and "doi:" in eloc:
                            doi = eloc.replace("doi: ","")
                        return {"title": rec["title"],
                            "authors": authors[:3],
                            "year": rec.get("pubdate","")[:4],
                            "source": "PubMed", "doi": doi,
                            "pmid": uid,
                            "url": f"https://pubmed.ncbi.nlm.nih.gov/{uid}/"}
    except Exception as e:
        sys.stderr.write(f"  [suggest] error: {e}\n")
    return None

# ── reference parsing ────────────────────────────────────

def parse_ref_text(text):
    refs = []
    lines = [l.strip() for l in text.strip().split('\n') if l.strip()]
    if lines and all(len(l) < 80 and (extract_doi(l) or extract_pmid(l) or l.startswith('10.')) for l in lines):
        for i, line in enumerate(lines):
            refs.append({"num": i+1, "raw": line, "doi": extract_doi(line), "pmid": extract_pmid(line)})
        return refs
    parts = re.split(r'\n\s*(?=\[?\d+\][.\s])', text.strip())
    if len(parts) < 2:
        parts = re.split(r'\n\s*(?=[A-Z][a-z]+[,\s])', text.strip())
    if len(parts) < 2:
        parts = [text.strip()]
    for i, part in enumerate(parts):
        part = part.strip()
        if not part: continue
        part = re.sub(r'^\[?\d+\]?[.\s]*', '', part)
        ref = {"num": i+1, "raw": part, "doi": extract_doi(part), "pmid": extract_pmid(part)}
        segs = part.split(". ")
        if len(segs) >= 2:
            ref["title"] = clean(segs[1]) if len(segs) > 1 else ""
            ref["authors"] = clean(segs[0])
            m = re.search(r'\b((?:19|20)\d{2})\b', part)
            ref["year"] = m.group(1) if m else None
        refs.append(ref)
    return refs

def extract_from_docx(path):
    if not INSTALLED["docx"]:
        return {"error": "python-docx not installed"}, [], []
    doc = Document(path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    ref_section = ""
    started = False
    for p in doc.paragraphs:
        t = p.text.strip().lower()
        if any(kw in t for kw in ["references", "bibliography", "works cited"]):
            started = True; continue
        if started:
            if len(t) > 20: ref_section += p.text + "\n"
            elif len(t) == 0 and len(ref_section) > 200: break
    if not ref_section:
        ref_section = full_text
    refs = parse_ref_text(ref_section)
    citations = re.findall(r'\([^)]*\d{4}[^)]*\)', full_text)
    unique = list(dict.fromkeys(citations))
    return {"title": doc.core_properties.title or ""}, refs, unique

# ── duplicates ───────────────────────────────────────────

def find_duplicates(refs):
    dupes = []
    for i, a in enumerate(refs):
        for j, b in enumerate(refs):
            if j <= i: continue
            score = 0
            if a.get("doi") and a["doi"] == b.get("doi"): score = 100
            elif a.get("year") and a["year"] == b.get("year"):
                aa = (a.get("authors","") or "")[:20].lower()
                bb = (b.get("authors","") or "")[:20].lower()
                if aa and bb and aa == bb: score = 90
            if score >= 80: dupes.append((i+1, j+1, score))
    return dupes

# ── citation matching ────────────────────────────────────

def match_citations(refs, citations):
    unmatched = []
    for cite in citations:
        m = re.search(r'([A-Z][a-z]+(?:\s+(?:et\s+al\.?|&\s+[A-Z][a-z]+))?)[,\s]+((?:19|20)\d{2})', cite)
        if not m:
            unmatched.append({"citation": cite, "reason": "Could not parse author+year"}); continue
        auth = m.group(1).split(",")[0].strip().lower().split()[-1]
        yr = m.group(2)
        if not any((ref.get("authors","") or "").split(",")[0].strip().lower() and auth in (ref.get("authors","") or "").split(",")[0].strip().lower() and ref.get("year") == yr for ref in refs):
            unmatched.append({"citation": cite, "reason": f"No reference entry for {auth} {yr}"})
    return unmatched

# ── report ───────────────────────────────────────────────

def format_report(meta, refs, dupes, unmatched, results, suggestions):
    lines = []
    lines.append("# Reference Audit Report")
    lines.append(f"**Date:** {time.strftime('%Y-%m-%d %H:%M')}")
    if meta.get("title"): lines.append(f"**Document:** {meta['title']}")
    lines.append(f"**Total references:** {len(refs)}\n")

    real = sum(1 for r in results if r.get("found"))
    fake = sum(1 for r in results if not r.get("found"))
    lines.append("## Summary")
    lines.append(f"- Confirmed real: **{real}**")
    lines.append(f"- Not found (suspicious): **{fake}**")
    lines.append(f"- Duplicate pairs: **{len(dupes)}**")
    lines.append(f"- Unmatched citations: **{len(unmatched)}**\n")

    suspicious = [r for r in results if not r.get("found")]
    if suspicious:
        lines.append("## Potentially Fake/Hallucinated References")
        lines.append("| # | Raw text | Suggested replacement |")
        lines.append("|---|----------|----------------------|")
        for r in suspicious:
            raw = r.get("raw","")[:60]
            sug = suggestions.get(r["num"])
            if sug:
                auth = ", ".join(sug["authors"][:2]) if isinstance(sug["authors"], list) else str(sug["authors"])
                lines.append(f"| {r['num']} | {raw}... | [{auth} ({sug['year']}) {sug['title'][:40]}]({sug['url']}) |")
            else:
                lines.append(f"| {r['num']} | {raw}... | No similar paper found |")
        lines.append("")

    if dupes:
        lines.append("## Duplicate References")
        lines.append("| Entry A | Entry B | Confidence |")
        lines.append("|---------|---------|------------|")
        for i,j,s in dupes: lines.append(f"| #{i} | #{j} | {s}% |")
        lines.append("")

    if unmatched:
        lines.append("## In-Text Citations Without Matching Reference")
        for u in unmatched: lines.append(f"- `{u['citation']}` -> {u['reason']}")
        lines.append("")

    lines.append("## All References")
    lines.append("| # | Reference | Status | Source | DOI/PMID |")
    lines.append("|---|-----------|--------|--------|----------|")
    for r in results:
        status = "OK" if r.get("found") else "??"
        src = r.get("source","") or "-"
        doi = r.get("doi","") or "-"
        raw = r.get("raw","")[:50]
        lines.append(f"| {r['num']} | {raw}... | {status} | {src} | {doi} |")
    lines.append("")
    return "\n".join(lines)

# ── main ─────────────────────────────────────────────────

def main():
    p = argparse.ArgumentParser(description="Verify references against real databases")
    sub = p.add_subparsers(dest="mode")
    pl = sub.add_parser("list"); pl.add_argument("--out")
    pd = sub.add_parser("docx"); pd.add_argument("path"); pd.add_argument("--out")
    pdoi = sub.add_parser("doi"); pdoi.add_argument("doi"); pdoi.add_argument("--out")
    ppm = sub.add_parser("pmid"); ppm.add_argument("pmid"); ppm.add_argument("--out")
    args = p.parse_args()
    if not args.mode: p.print_help(); return

    meta, refs = {}, []
    if args.mode == "doi": refs = [{"num":1,"raw":args.doi,"doi":args.doi}]
    elif args.mode == "pmid": refs = [{"num":1,"raw":args.pmid,"pmid":args.pmid}]
    elif args.mode == "list": refs = parse_ref_text(sys.stdin.read())
    elif args.mode == "docx":
        meta, refs, citations = extract_from_docx(args.path)
        if "error" in meta: print(json.dumps(meta)); return

    sys.stderr.write(f"Verifying {len(refs)} references...\n")
    results = []
    for ref in refs:
        result = verify_ref(ref)
        results.append(result)
        mark = "OK" if result.get("found") else "??"
        raw = (result.get("raw","") or "")[:50].replace('\n',' ')
        sys.stderr.write(f"  [{result['num']}] {mark} {raw}\n")
        time.sleep(0.3)

    sys.stderr.write("Searching for replacement suggestions...\n")
    suggestions = {}
    for r in results:
        if not r.get("found"):
            sug = suggest_replacement(r.get("raw",""))
            if sug:
                suggestions[r["num"]] = sug
                sys.stderr.write(f"  [{r['num']}] Suggestion: {sug['title'][:50]}\n")
            else: sys.stderr.write(f"  [{r['num']}] No suggestion\n")

    dupes = find_duplicates(refs)
    unmatched = match_citations(refs, citations) if args.mode == "docx" else []
    report = format_report(meta, refs, dupes, unmatched, results, suggestions)

    out = args.out or "refcheck_report.md"
    with open(out, "w", encoding="utf-8") as f: f.write(report)
    print(f"\nReport saved to: {out}")
    print(report)

if __name__ == "__main__":
    main()
